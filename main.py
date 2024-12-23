import unittest

from docx import Document
from docx.shared import Inches
import os
from unittest.mock import patch, MagicMock
import re

FILENAME = "demo.docx"
TEMPFILENAME = "temp.docx"

class FoodOrder:
    def __init__(self, name, foodName, amount, status, imgPath):
        self.name = name
        self.foodName = foodName
        self.amount = amount
        self.status = status
        self.imgPath = imgPath
    
    def addFood(self):
        document = Document(FILENAME)
        document.add_paragraph(
            f"Nama : {self.name}, Menu : {self.foodName}, Jumlah : {self.amount}, Status : {self.status}"
        )
        document.add_picture(self.imgPath, width=Inches(2.0))
        document.save(FILENAME)

def checkDocument():
    if not os.path.exists(FILENAME):
        document = Document()
        document.add_heading("Manajemen Makanan", level=1)
        document.save(FILENAME)


# def updateStatus(data, newStatus):
#     document = Document(FILENAME)

#     for paragraph in document.paragraphs:
#         if "Nama :" in paragraph.text:
#             # Pisahkan teks berdasarkan koma
#             parts = paragraph.text.split(", ")
            
#             # Ambil bagian nama
#             numPart = parts[0].split("_")[1].strip()  # Mengambil 'Person_1' dari 'Nama : Person_1'
            
#             targetPart = data.split("_")[1].strip()
            
#             if int(numPart) == int(targetPart):  # Cocokkan nama dengan parameter
#                 # Cari dan ubah bagian 'Status'
#                 for i, part in enumerate(parts):
#                     if "Status :" in part:
#                         parts[i] = f"Status : {newStatus}"  # Update status
#                         break
                
#                 # Gabungkan kembali teks paragraf
#                 paragraph.text = ", ".join(parts)
#                 break  # Setelah menemukan dan memperbarui, keluar dari loop

#     # Simpan perubahan ke file
#     document.save(FILENAME)


# def updateStatus(name, newStatus):
#     document = Document(FILENAME)
#     isUpdated = False  

#     for paragraph in document.paragraphs:
#         if name in paragraph.text:
#             if "Status :" in paragraph.text:
#                 if name == paragraph:
#                     print(f"Old text is {paragraph.text}")
#                     oldText = paragraph.text
#                     newText = oldText.split("Status :")[0] + f"Status : {newStatus}"
#                     paragraph.text = newText
#                     isUpdated = True
#                     print(f"new text is {paragraph.text}\n\n")

#     if isUpdated:
#         document.save(FILENAME)
#         # print(f"Status untuk {name} berhasil diubah.")
#     else:
#          print(f"Status untuk {name} gagal diubah.")
def updateStatus(name, newStatus):
    document = Document(FILENAME)
    isUpdated = False  

    for paragraph in document.paragraphs:
        if "Status" in paragraph.text:
            parts = paragraph.text.split(",")
            newParts = parts[0].split(" ")[2:]
            newPartsString = " ".join(newParts)

            # print(f"{newPartsString}")
            # print(f"{name}")
            # print("\n")
            if name == newPartsString:
                print(f"Old text is {paragraph.text}")
                oldText = paragraph.text
                newText = oldText.split("Status :")[0] + f"Status : {newStatus}"
                paragraph.text = newText
                isUpdated = True
                # print(f"new text is {paragraph.text}\n\n")
                break

    if isUpdated:
        document.save(FILENAME)
        print(f"Status untuk {name} berhasil diubah.")
    else:
         print(f"Status untuk {name} gagal diubah.")

def removeFoodOrder(target):
    document = Document(FILENAME)
    arrFoodOrder = []

    isFound = False

    for paragraph in document.paragraphs:
        if "Nama :" in paragraph.text and "Menu :" in paragraph.text:
            data = paragraph.text.split(", ")
            name = data[0].split(":")[1].strip()
            foodName = data[1].split(":")[1].strip()
            amount = data[2].split(":")[1].strip()
            status = data[3].split(":")[1].strip()

            imgPath = "C:\\Users\\USER\\Downloads\\img\\" + foodName.lower() + ".jpg"  

            if target in name and status.lower() == "batal":
                isFound = True
            else:
                arrFoodOrder.append(FoodOrder(name, foodName, amount, status, imgPath))

    if isFound:
        tempDocument = Document()
        tempDocument.add_heading("Manajemen Makanan", level=1)

        for order in arrFoodOrder:
            tempDocument.add_paragraph(
                f"Nama : {order.name}, Menu : {order.foodName}, Jumlah : {order.amount}, Status : {order.status}"
            )
            tempDocument.add_picture(order.imgPath, width=Inches(2.0))

        tempDocument.save(TEMPFILENAME)

        os.remove(FILENAME)  
        os.rename(TEMPFILENAME, FILENAME) 
        print(f"Pesanan milik {target} dengan status 'Batal' telah dihapus. Dokumen diperbarui.")
    else:
        print(f"Tidak ditemukan pesanan milik {target} dengan status 'Batal'. Tidak ada perubahan pada dokumen.")


def showOrders():
    document = Document(FILENAME)
    print("\nPesanan yang ada:")
    
    order_number = 1

    for paragraph in document.paragraphs:
        if "Nama :" in paragraph.text and "Menu :" in paragraph.text:
            print(f"{order_number}. {paragraph.text}\n")  
            order_number +=1

def searchOrderByName(name):
    document = Document(FILENAME)
    found = False  
    for paragraph in document.paragraphs:
        if "Nama :" in paragraph.text and name.lower() in paragraph.text.lower():
            print(paragraph.text)  
            found = True
            break
    if not found:
        print(f"Tidak ada pesanan ditemukan untuk {name}.")



# checkDocument()

while False:
        print("\nManajemen Pesanan Restoran")
        print("1. Tambah Pesanan")
        print("2. Tampilkan Pesanan")
        print("3. Update Pesanan")
        print("4. Hapus Pesanan")
        print("5. Cari Pesanan")
        print("6. Keluar")
        
        pilihan = input("Pilih menu (1-5): ")

        if pilihan == "1":
            print("\nTambah Pesanan Baru")
            name = input("Masukkan Nama Pelanggan : ")
            foodName = input("Masukkan Menu yang dipesan (takoyaki,ramen,sushi): ")
            amount = input("Masukkan Jumlah Pesanan: ")
            imgPath = input("Masukkan Path Gambar : ")
            if os.path.exists(imgPath):
                defaultStatus = "Diposes"
                order = FoodOrder(name, foodName, amount,defaultStatus,imgPath)
                order.addFood()
                print(f"Pesanan {foodName} oleh {name} berhasil ditambah.")
            else:
                print("Gambar tidak ditemukan, silahkan coba lagi")
        elif pilihan == "2":
            showOrders()
        elif pilihan == "3":
            print("Data Pesanan :\nManajemen Pesanan Restoran")
            showOrders()
            target = input("Masukkan nama pelanggan yang ingin diupdate : ")
            newStatus = input("Masukkan status baru (Proses/Selesai/Batal) : ")
            updateStatus(target, newStatus)

        elif pilihan == "4":
            print("Data Pesanan :")
            showOrders()
            target = input("Masukkan nama pelanggan yang ingin dihapus pesannya: ")
            removeFoodOrder(target)

        elif pilihan == "5":
            target = input("Masukkan nama pelanggan yang ingin dicari : ")
            searchOrderByName(target)
        elif pilihan == "6":
            print("Terima kasih! Program selesai.")
            break
        else:
            print("Pilihan tidak valid. Silakan pilih antara 1-5.")
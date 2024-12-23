import unittest
from docx import Document
import os
import random
from main import FoodOrder  

FILENAME = "demo.docx"

class TestFoodOrder(unittest.TestCase): 
    def testAddFood(self):
        # Daftar makanan yang akan diuji
        foodNames = ["Ramen", "Takoyaki", "Sushi"]
        
        imgPaths = ["ramen.jpg", "takoyaki.jpg", "sushi.jpg"]
        status = "Diproses"
        foodOrders = []
        
        # Generate 100 data
        for i in range(1, 101):
            name = f"Person_{i}"
            menu = random.choice(foodNames)  
            amount = random.randint(1, 9)  
            imgFile = imgPaths[foodNames.index(menu)] 
            foodOrders.append(FoodOrder(name, menu, str(amount), status, imgFile))
        
        # String yang diharapkan
        expectedTexts = [
            f"Nama : {foodOrder.name}, Menu : {foodOrder.foodName}, Jumlah : {foodOrder.amount}, Status : {foodOrder.status}"
            for foodOrder in foodOrders
        ]
        
        # Validasi apakah file tersedia, jika tidak akan membuat file baru dan heading
        checkDocument()

        failureCountStr = 0  # Menghitung kegagalan menemukan teks
        imgFoodType = 0  # Menghitung jenis gambar makanan

        for foodOrder, expectedText in zip(foodOrders, expectedTexts):
            # Menjalankan fungsi addFood untuk menambahkan informasi makanan ke dalam demo.docx
            foodOrder.addFood()

            doc = Document(FILENAME)
            
            # Memeriksa apakah teks makanan ada dalam dokumen
            isFoundText = False
            for paragraph in doc.paragraphs:
                cleanedText = paragraph.text.strip()
                if expectedText in cleanedText:
                    self.assertIn(expectedText, cleanedText)
                    isFoundText = True
            if not isFoundText:
                print(f"Teks '{expectedText}' yang diharapkan tidak ditemukan dalam dokumen.")
                failureCountStr += 1
                isFoundText = False
            
        # Memeriksa apakah gambar yang terkait dengan makanan ada dalam dokumen
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref or "media" in rel.target_ref: 
                imgFoodType += 1
        
        self.assertEqual(failureCountStr, 0, f"Terdapat data yang tidak ditemukan berjumlah {failureCountStr}")
        self.assertEqual(imgFoodType, 3, f"Gambar yang ditemukan tidak sesuai nama gambarnya, ditemukan {3 - imgFoodType} gambar.")
        
        if failureCountStr == 0 and imgFoodType == 3:
            print(f"=== Semua teks ditemukan dan jumlah gambar 100 ===")

# Inisialisasi awal file
def checkDocument():
    if not os.path.exists(FILENAME):
        document = Document()
        document.add_heading("Manajemen Makanan", level=1)
        document.save(FILENAME)

if __name__ == "__main__":
    unittest.main()

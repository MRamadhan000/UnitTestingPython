import unittest
from docx import Document
import os
from main import removeFoodOrder

FILENAME = "demo.docx"

class TestFoodOrder(unittest.TestCase): 
    def testDeleteAllData(self):
    # Daftar nama untuk dihapus
        namesArr = [f"Person_{i}" for i in range(1, 101)]  # Person_1 hingga Person_100

        # Memanggil removeFoodOrder untuk setiap nama dalam daftar
        for name in namesArr:
            removeFoodOrder(name)

        # Membaca ulang dokumen untuk validasi
        doc = Document(FILENAME)
        for name in namesArr:
            found = False
            for paragraph in doc.paragraphs:
                if name in paragraph.text:
                    if "Status : Batal" in paragraph.text:
                        found = True
                        break
            # Pastikan nama dengan status "Batal" tidak ditemukan
            self.assertFalse(found, f"{name} masih ditemukan dengan status Batal di dokumen")
            print(f"{name} berhasil dihapus dari dokumen")

# Inisialisasi awal file
def checkDocument():
    if not os.path.exists(FILENAME):
        document = Document()
        document.add_heading("Manajemen Makanan", level=1)
        document.save(FILENAME)

if __name__ == "__main__":
    unittest.main()
